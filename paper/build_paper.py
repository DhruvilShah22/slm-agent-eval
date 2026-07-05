"""Builds paper/paper.docx — the arXiv submission manuscript.

Regenerable: python paper/build_paper.py. Numbers in the text were taken from
analysis/results.json (commit-pinned); figures are embedded from
analysis/figures/. Single-column, Times New Roman, standard preprint layout.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
FIGS = ROOT / "analysis" / "figures"
OUT = ROOT / "paper" / "paper.docx"

AUTHOR = "Dhruvil Shah"
EMAIL = "[EMAIL — confirm before submission]"
AFFIL = "Independent Researcher"

TITLE = ("Quantization, Not Just Scale: Reliability and Failure Anatomy of "
         "Small Open-Weight Tool Agents")

ABSTRACT = (
    "Small open-weight language models (1–7B parameters) are the only "
    "agents most practitioners can run at zero cost, yet agent reliability "
    "has mostly been measured on frontier APIs. We build a five-tool agent "
    "harness with a 25-task suite spanning five probe slices (single-tool "
    "use, multi-tool chains, distractor tools, underspecified requests, and "
    "injected tool faults) and run 2,400 seeded episodes across a model-size "
    "× quantization × mitigation matrix, graded programmatically "
    "(blind-validated, κ = 0.86 for success and 0.94 for failure "
    "category). Three results stand out. First, quantization dominates: at "
    "1.5B, the 8-bit build more than doubles task success over the default "
    "4-bit download (31.0% vs 15.0%; +16 points, Holm-corrected exact "
    "McNemar p = 2×10⁻⁵), matching the gain from doubling "
    "parameters, and twice in our matrix a higher-precision smaller model "
    "matched a larger 4-bit one. The damage is qualitative, not just "
    "quantitative: the 4-bit model often never calls a tool, while the 8-bit "
    "model calls tools and fails later, at synthesis. Second, a "
    "deterministic verification-and-retry guardrail (schema validation, "
    "typed errors, bounded retries) never significantly improves overall "
    "success, and significantly reduces it in two different model families "
    "(−7.5 points at Qwen2.5-1.5B-Q8, −6.5 at Llama-3.2-1B); most "
    "failures originate at tool selection or answer synthesis, upstream of "
    "anything an argument-level check can reach. The same guardrail, "
    "however, triples recovery from injected tool faults at 3B, indicating "
    "the mechanism works once a model can act on error feedback. Third, "
    "reliability collapses under repetition: the best cell completes all "
    "eight seeded repeats of a task only 16% of the time, and 1.5B models "
    "recovered from 0 of 84 injected faults, typically hallucinating a "
    "concrete answer instead. The harness, raw logs, and analysis are "
    "public, every reported number regenerates from the logs in CI, and the "
    "full study reproduces in under four GPU-hours on a free tier or on a "
    "consumer CPU.")

FIG_CAPTIONS = {
    "fig1_success_rates.png":
        "Figure 1: Task success rate by cell with Wilson 95% intervals "
        "(core matrix, n = 200 episodes per cell).",
    "fig2_passk.png":
        "Figure 2: pass^k — the probability that all k seeded repeats "
        "of a task succeed — for the six core cells.",
    "fig3_first_failures.png":
        "Figure 3: First-failure composition of all 200 episodes per core "
        "cell. Quantization changes how the 1.5B model fails, not only how "
        "often: no_tool_call dominates at Q4 and nearly vanishes at Q8.",
    "fig4_fault_recovery.png":
        "Figure 4: Outcomes of injected tool faults (slice S5, core matrix). "
        "Recovery appears almost exclusively in the 3B guardrail cell.",
}

REFERENCES = [
    "P. Belcak, G. Heinrich, S. Diao, Y. Fu, X. Dong, S. Muralidharan, "
    "Y. C. Lin, and P. Molchanov. Small language models are the future of "
    "agentic AI. arXiv:2506.02153, 2025.",
    "S. Yao, N. Shinn, P. Razavi, and K. Narasimhan. τ-bench: A "
    "benchmark for tool-agent-user interaction in real-world domains. "
    "arXiv:2406.12045, 2024.",
    "X. Liu et al. AgentBench: Evaluating LLMs as agents. In Proc. ICLR, "
    "2024. arXiv:2308.03688.",
    "Berkeley Function Calling Leaderboard. "
    "https://gorilla.cs.berkeley.edu/leaderboard.html, 2024–2026.",
    "R. S. Babu and A. Agrawal. Self-healing agentic orchestrators for "
    "reliable tool-augmented large language model systems. "
    "arXiv:2606.01416, 2026.",
    "P. Dong, Z. Tang, X. Liu, L. Li, X. Chu, and B. Li. Can compressed "
    "LLMs truly act? An empirical evaluation of agentic capabilities in "
    "LLM compression. arXiv:2505.19433, 2025.",
    "M. A. Haque, F. Rahman, K. D. Gupta, K. Shujaee, and R. George. "
    "TinyLLM: Evaluation and optimization of small language models for "
    "agentic tasks on edge devices. arXiv:2511.22138, 2025.",
    "D. Zhu, X. Ma, Y. Shen, X. Li, Y. Zhao, S. Wang, L. Yan, and D. Yin. "
    "When tools fail: Benchmarking dynamic replanning and anomaly recovery "
    "in LLM agents. arXiv:2606.05806, 2026.",
    "W. Yang, C. Song, X. Li, D. Ganguly, C. Ma, S. Wang, Z. Dou, Y. Zhou, "
    "V. Chaudhary, and X. Han. AgentCE-Bench: Agent configurable evaluation "
    "with scalable horizons and controllable difficulty under lightweight "
    "environments. arXiv:2604.06111, 2026.",
    "S. V. Vuddanti, A. Shah, S. K. Chittiprolu, T. Song, S. Dev, K. Zhu, "
    "and M. Chaudhary. PALADIN: Self-correcting language model agents to "
    "cure tool-failure cases. arXiv:2509.25238, 2025.",
    "M. Cemri et al. Why do multi-agent LLM systems fail? In NeurIPS "
    "Datasets and Benchmarks, 2025. arXiv:2503.13657.",
    "Y. Ruan, H. Dong, A. Wang, S. Pitis, Y. Zhou, J. Ba, Y. Dubois, "
    "C. J. Maddison, and T. Hashimoto. Identifying the risks of LM agents "
    "with an LM-emulated sandbox. In Proc. ICLR, 2024. arXiv:2309.15817.",
    "N. F. Liu, K. Lin, J. Hewitt, A. Paranjape, M. Bevilacqua, F. Petroni, "
    "and P. Liang. Lost in the middle: How language models use long "
    "contexts. TACL, 2024.",
    "W. Zhang et al. Learning to ask: When LLM agents meet unclear "
    "instruction. In Proc. EMNLP, 2025.",
]


def body(doc: Document, text: str, first_indent: bool = True) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(0.2)


def heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13 if level == 1 else 11.5)


def table(doc: Document, header: list[str], rows: list[list[str]],
          caption: str) -> None:
    cap = doc.add_paragraph(caption)
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True
    cap.paragraph_format.space_before = Pt(8)
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for j, htxt in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = htxt
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def figure(doc: Document, filename: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIGS / filename), width=Inches(5.6))
    cap = doc.add_paragraph(FIG_CAPTIONS[filename])
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1.1)
        section.top_margin = section.bottom_margin = Inches(1.0)

    # ---- Title block ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(TITLE)
    run.font.size = Pt(16)
    run.font.bold = True
    a = doc.add_paragraph()
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a.add_run(f"{AUTHOR}\n").font.size = Pt(11.5)
    a.add_run(f"{AFFIL}\n{EMAIL}").font.size = Pt(10)

    # ---- Abstract ----
    ah = doc.add_paragraph()
    ah.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = ah.add_run("Abstract")
    ar.font.bold = True
    ar.font.size = Pt(11)
    ab = doc.add_paragraph(ABSTRACT)
    ab.paragraph_format.left_indent = Inches(0.35)
    ab.paragraph_format.right_indent = Inches(0.35)
    ab.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in ab.runs:
        run.font.size = Pt(9.5)

    # ---- 1 Introduction ----
    heading(doc, "1  Introduction", 1)
    body(doc, (
        "If you want a language-model agent and cannot pay per token, your "
        "options narrow quickly: a 1–7B open-weight model, quantized to "
        "fit consumer hardware, served locally. A growing position in the "
        "literature holds that such small models are the economically "
        "rational substrate for agentic systems [1]. Whether they are "
        "*reliable* enough is a different question, and the evidence so far "
        "comes mostly from the other end of the spectrum: τ-bench "
        "showed that even GPT-4o solves fewer than half of its tasks and is "
        "inconsistent across repeated trials [2]; AgentBench found "
        "open-weight models trailing API models by large margins [3]; "
        "function-calling leaderboards rank models by single-shot accuracy "
        "[4]. None of this tells a practitioner what to expect from the "
        "quantized 1.5B model they just pulled, how it fails, or whether the "
        "obvious cheap fix is worth the latency."), first_indent=False)
    body(doc, (
        "We ask four questions about small open-weight tool agents. How "
        "reliably do they complete multi-step tool tasks, measured across "
        "repeated seeded trials rather than single shots? Where in the "
        "pipeline do failures start — tool selection, argument "
        "construction, error handling, or answer synthesis? Does a "
        "deterministic verification-and-retry guardrail close any of the "
        "gap, and at what cost? And how does weight quantization — the "
        "knob every local deployment turns — modulate all of the "
        "above?"))
    body(doc, (
        "We answer these with a deliberately transparent testbed: a "
        "five-tool customer-support agent over a fully synthetic world, a "
        "25-task suite with five probe slices, programmatic grading "
        "validated blind against hand labels, and 2,400 episodes across "
        "five model×quantization settings, each with and without the "
        "guardrail. Hypotheses, contrasts, and gates were fixed in a design "
        "document before any experiment ran; one of the four hypotheses was "
        "falsified and is reported as such. Our contributions:"))
    for bullet in (
        "A reliability-first measurement of 1–7B GGUF-quantized agents "
        "with pass^k, per-stage failure attribution, and paired statistics "
        "— the regime prior agent benchmarks mostly skip.",
        "Evidence that quantization is a first-order variable for agentic "
        "use: +16 points from Q4_K_M to Q8_0 at 1.5B (p = 2×10⁻"
        "⁵), with a qualitative shift in failure mode, and two "
        "instances where a smaller higher-precision model matched a larger "
        "4-bit one. Static compression benchmarks predicted 1–3% [6].",
        "A cautionary result for the most common mitigation: schema "
        "validation with typed-error retries never significantly helped "
        "overall success in five settings and significantly hurt in two "
        "model families — while still tripling fault recovery at 3B, "
        "locating a capability threshold below which verification feedback "
        "is noise.",
        "A zero-cost, fully reproducible harness: every number in this "
        "paper is regenerated from committed raw logs by CI, and the whole "
        "study reruns in under four GPU-hours on a free tier."):
        p = doc.add_paragraph(bullet, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10.5)

    # ---- 2 Related work ----
    heading(doc, "2  Related work", 1)
    body(doc, (
        "Agent benchmarks and reliability. τ-bench introduced pass^k "
        "over repeated trials and an LLM-simulated user, applied to "
        "frontier API models [2]. We adopt pass^k and move the object of "
        "study to small local models; our clarification turns are scripted "
        "rather than LLM-simulated, which keeps the study free and "
        "deterministic. AgentBench established the API-vs-open gap across "
        "eight environments [3]. BFCL defines the standard schema-level "
        "grading of function calls [4], which our grader follows, but "
        "leaderboard accuracy is single-shot and unattributed; we add "
        "repetition, attribution, and a mitigation arm."), first_indent=False)
    body(doc, (
        "Mitigations and self-healing. Babu and Agrawal treat reliability "
        "as a bounded runtime-control problem and report 98.8% success on a "
        "100-task fault-injection benchmark with budgeted recovery and "
        "verification [5]. Their evaluation abstracts the model away "
        "(unnamed, no repeated-trial metrics or intervals, call-count cost "
        "proxies), which is precisely the gap we fill with named quantized "
        "models, paired tests, and wall-clock accounting — and our "
        "results add a boundary condition their abstraction cannot see: "
        "below a capability threshold, verification feedback reduces "
        "success. PALADIN trains recovery behavior from failure-injected "
        "trajectories [10]; our guardrail is deliberately training-free. "
        "Tool-fault benchmarks (ToolMaze [8], AgentCE-Bench [9]) inject "
        "perturbations at larger scale than our S5 slice; we borrow the "
        "idea and add the small-model recovery numbers."))
    body(doc, (
        "Compression and agentic ability. ACBench evaluated GPTQ/AWQ "
        "compression on agent-relevant benchmarks and found 4-bit "
        "quantization costs only 1–3% on tool use [6]. Our results "
        "disagree in the live-loop setting by an order of magnitude at "
        "1.5B, suggesting single-shot static evaluation understates "
        "quantization damage to sequential agentic behavior. TinyLLM "
        "studies sub-3B function calling and improves it by fine-tuning "
        "[7]; we measure, rather than train, and target the inference-time "
        "question. MAST taxonomizes failures of multi-agent systems [11]; "
        "our taxonomy is single-agent and pipeline-staged, designed to be "
        "assignable by deterministic rules. ToolEmu uses an LM-emulated "
        "sandbox for safety risks [12]. Clarification under ambiguity is an "
        "active method area [14]; our S4 slice only measures the behavior. "
        "Long-context degradation [13] is out of scope by design."))

    # ---- 3 Testbed ----
    heading(doc, "3  Testbed", 1)
    body(doc, (
        "Agent loop. A minimal ReAct-style loop with native tool calling "
        "through Ollama's chat API — no agent framework, so every "
        "token the model sees is in the logs. System prompt, task, then "
        "up to ten model calls; temperature 0.7, top-p 0.9, context 4,096 "
        "tokens, seeds 1–8 per task. Episodes are logged in full "
        "(requests, responses, tool executions, timing) and written "
        "atomically, so runs resume losslessly."), first_indent=False)
    body(doc, (
        "Tools. Five local, deterministic tools: an AST-whitelisted "
        "calculator; BM25 search over a 64-document policy corpus; two "
        "parameterized SQLite lookups (get_order, find_products) over a "
        "seeded 40-product, 30-order database; and a sandboxed Python "
        "executor. The world (“Zephyra Outfitters”) is entirely "
        "fictional and generated from a fixed seed, so no answer can come "
        "from memorization: an ungrounded answer is detectably wrong."))
    body(doc, (
        "Tasks. Twenty-five hand-written tasks in five slices: S1 "
        "single-tool (5), S2 multi-tool chains (8), S3 distractor — a "
        "plausible but wrong tool is available (4), S4 underspecified "
        "requests, where the harness answers one clarifying question with "
        "a scripted reply (4), and S5 injected tool faults (error, "
        "timeout, or empty result on the first call, clearing afterwards) "
        "(4). Gold answers are resolved from the generated world at "
        "grading time rather than stored, so data and answers cannot "
        "drift; S4 golds are conditional on whether the agent asked."))
    body(doc, (
        "Grading and failure attribution. Success is a normalized match "
        "of the extracted final answer (numeric tolerance, ISO date, or "
        "fuzzy string). Each failed episode receives a first-failure "
        "label by deterministic rules over its event log: no_tool_call, "
        "wrong_tool, malformed_args, bad_arg_values, ignored_tool_error, "
        "synthesis_error, or max_turns; S4 episodes are additionally coded "
        "asked / looked-up / guessed, and S5 episodes carry recovery "
        "flags. No LLM judge is used anywhere. To validate the grader, 42 "
        "stratified episodes were hand-labeled blind (trajectories printed "
        "without classifier output): agreement was 40/42 on success "
        "(κ = 0.86) and 40/42 on failure category (κ = 0.94); "
        "both disagreements traced to one string-normalization gap (a "
        "contraction), fixed with an edit-distance rule, after which the "
        "same sample agrees 42/42 (in-sample). Two earlier grader defects "
        "caught in the pilot are documented in the repository log."))
    body(doc, (
        "Guardrail. The mitigation under test is deterministic and "
        "model-free: every tool call is validated against its JSON schema "
        "before execution; violations return a typed error naming the "
        "exact problem, with at most two consecutive retries per tool and "
        "four extra model calls per episode; tool runtime errors are "
        "wrapped as structured, retriable error objects. The baseline "
        "passes calls through unchanged and lets the tool layer fail the "
        "way a real API does — terse strings. Schema validity is "
        "logged in both conditions, so malformed-call rates are always "
        "measurable."))

    # ---- 4 Experimental design ----
    heading(doc, "4  Experimental design", 1)
    body(doc, (
        "Core matrix: Qwen2.5-Instruct at 1.5B-Q4_K_M, 1.5B-Q8_0, and "
        "3B-Q4_K_M (GGUF builds as shipped by Ollama; digests pinned in "
        "run manifests), each baseline and guardrail: 6 cells × 25 "
        "tasks × 8 seeds = 1,200 episodes. Extension matrix "
        "(exploratory, same protocol): Llama-3.2-1B (Q8_0), "
        "Qwen2.5-3B-Q8_0, and Qwen2.5-7B-Q4_K_M, again both conditions: "
        "1,200 further episodes. Episodes ran on a free Kaggle P100 in "
        "56 minutes per 1,200; the identical harness runs on a consumer "
        "CPU (throughput measured and reported in the repository)."),
        first_indent=False)
    body(doc, (
        "Statistics. Five contrasts were pre-registered and paired on "
        "(task, seed): the guardrail effect within each core "
        "model×quantization, Q8 vs Q4 at 1.5B, and 3B vs 1.5B at Q4 "
        "— exact McNemar tests with Holm correction, risk differences "
        "with cluster-bootstrap 95% intervals (10,000 resamples, clustered "
        "by task). Per-cell uncertainty uses Wilson intervals; reliability "
        "uses the unbiased pass^k estimator over 8 seeds, averaged over "
        "tasks. Analysis re-grades every episode from raw logs; grades "
        "stored at run time are advisory."))

    # ---- 5 Results ----
    heading(doc, "5  Results", 1)
    heading(doc, "5.1  No cell is reliable, and repetition is brutal", 2)
    table(doc,
          ["Cell", "Model", "Condition", "Success", "95% CI", "pass^8"],
          [["C1", "1.5B Q4_K_M", "baseline", "15.0%", "10.7–20.6", "0.00"],
           ["C2", "1.5B Q4_K_M", "guardrail", "12.5%", "8.6–17.8", "0.00"],
           ["C3", "1.5B Q8_0", "baseline", "31.0%", "25.0–37.7", "0.04"],
           ["C4", "1.5B Q8_0", "guardrail", "23.5%", "18.2–29.8", "0.00"],
           ["C5", "3B Q4_K_M", "baseline", "30.5%", "24.5–37.2", "0.08"],
           ["C6", "3B Q4_K_M", "guardrail", "33.5%", "27.3–40.3", "0.08"]],
          "Table 1: Core matrix (200 episodes per cell). Wilson intervals; "
          "pass^8 is the estimated probability all eight seeds succeed.")
    body(doc, (
        "Success rates sit between 12.5% and 33.5% (Table 1, Figure 1), "
        "and the decay under repetition is steep everywhere (Figure 2): "
        "the best core cell would survive eight repeats of a task 8% of "
        "the time. This extends the frontier-model inconsistency of [2] "
        "into the small-model regime, where it is far more severe "
        "(H1 supported)."), first_indent=False)
    figure(doc, "fig1_success_rates.png")
    figure(doc, "fig2_passk.png")

    heading(doc, "5.2  Quantization is a first-order variable", 2)
    body(doc, (
        "Q8_0 versus Q4_K_M at 1.5B (baseline): +16.0 points (cluster-"
        "bootstrap CI +6.0 to +27.0; discordant pairs 40:8; Holm "
        "p = 2×10⁻⁵) — the same magnitude as doubling "
        "parameters at fixed quantization (3B-Q4 vs 1.5B-Q4: +15.5, CI "
        "+0.5 to +31.0, p = 3.5×10⁻⁴). Figure 3 shows the "
        "effect is qualitative: the Q4 build's modal first failure is "
        "no_tool_call (61/200 — it answers from nothing, often "
        "emitting pseudo-JSON), which nearly vanishes at Q8 (5/200); Q8's "
        "failures migrate downstream to synthesis. On ambiguity probes "
        "the Q8 build asks a clarifying question twice as often (44% vs "
        "22%). The pattern recurs in the extension matrix at 3B "
        "(Q8 43.5% vs Q4 30.5%, descriptive), and twice a smaller "
        "higher-precision model matched a larger 4-bit one (1.5B-Q8 ≈ "
        "3B-Q4; 3B-Q8 ≈ 7B-Q4). Under a fixed memory budget, "
        "precision appears to buy as much agentic reliability as "
        "parameters — static compression benchmarks [6] do not "
        "predict this (H3 supported at 1.5B)."), first_indent=False)
    figure(doc, "fig3_first_failures.png")

    heading(doc, "5.3  The guardrail never helps overall — and can hurt "
                 "(H2 falsified)", 2)
    table(doc,
          ["Contrast", "RD", "95% CI", "Discordants", "Holm p"],
          [["guardrail @1.5B-Q4", "−2.5", "−9.5, +3.0", "3:8", "0.42"],
           ["guardrail @1.5B-Q8", "−7.5", "−17.5, −0.5",
            "1:16", "8×10⁻⁴"],
           ["guardrail @3B-Q4", "+3.0", "−4.0, +11.0", "11:5", "0.42"],
           ["guardrail @1B (Llama, ext.)", "−6.5", "−15.0, 0.0",
            "2:15", "7×10⁻³"],
           ["guardrail @3B-Q8 (ext.)", "−1.5", "−4.0, +0.5",
            "4:7", "1.0"],
           ["guardrail @7B-Q4 (ext.)", "−0.5", "−1.5, 0.0",
            "0:1", "1.0"]],
          "Table 2: Guardrail effect (risk difference, percentage points), "
          "paired on (task, seed). Core rows Holm-corrected within the "
          "pre-registered family; extension rows corrected within theirs.")
    body(doc, (
        "In five model×quantization settings the guardrail never "
        "significantly improved success, and in two — from two "
        "different model families — it significantly reduced it "
        "(Table 2). At 1.5B-Q8 it flipped sixteen task-seed pairs from "
        "success to failure while rescuing one. The logs show why. First, "
        "most failures begin upstream of arguments (Figure 3): wrong or "
        "missing tool selection and bad synthesis are inert to schema "
        "validation. Second, when the guardrail does block a malformed "
        "call, the verbose typed error often derails a small model rather "
        "than repairing it — one Llama episode re-sent an invented "
        "order_by argument eight times, alternating blocked and "
        "passed-through, until the turn budget expired. Malformed-args as "
        "a first failure rises under the guardrail at 1.5B-Q8 (19 → "
        "34) because retries multiply chances to fail at the same step. "
        "Llama-3.2-1B replicates the harm with a different fingerprint: "
        "its dominant failure is inventing argument names (malformed_args "
        "86–97/200), and the guardrail still costs it 6.5 points."),
        first_indent=False)

    heading(doc, "5.4  Error recovery is a capability cliff — and the "
                 "guardrail's one real win", 2)
    body(doc, (
        "Across all four 1.5B cells, zero of 84 fired faults were "
        "recovered — not one retry. Afterwards the models usually "
        "answered anyway: 22 of 25 fired faults in the 1.5B-Q4 baseline "
        "ended in a fabricated concrete answer (a status, a price), "
        "versus three honest refusals. Typed errors shift some "
        "fabrications to refusals (22 → 13) without enabling "
        "recovery — arguably the guardrail's only benefit at that "
        "scale, and a safety-relevant one. At 3B-Q4 the guardrail's "
        "structured, retriable errors raise recovery from 3/16 to 10/17 "
        "(Figure 4), and ignored_tool_error as a first failure halves. "
        "The extension matrix explains the mechanism: native recovery "
        "tracks model quality — 3B-Q8 recovers 15/24 with no "
        "guardrail at all, 7B-Q4 8/23 — so the guardrail at 3B-Q4 "
        "was compensating for what Q4 quantization had removed. "
        "Verification feedback, in short, is only useful to a model "
        "capable of acting on it."), first_indent=False)
    figure(doc, "fig4_fault_recovery.png")

    heading(doc, "5.5  Cost", 2)
    body(doc, (
        "The guardrail's paired overhead is small: +0.16 s and +322 "
        "tokens per episode at 3B (+3.7% wall-clock), approximately zero "
        "at 1.5B (H4 supported). Derived consumer-CPU episode times "
        "(token counts × measured laptop throughput; derived, not "
        "measured) are ≈50 s at 1.5B-Q4 and ≈225–245 s at "
        "3B-Q4. The efficiency question is thus not the guardrail's "
        "latency but whether it does anything useful — below 3B, it "
        "does not."), first_indent=False)

    # ---- 6 Discussion ----
    heading(doc, "6  Discussion", 1)
    body(doc, (
        "Three practical readings. For practitioners: the default 4-bit "
        "download is the wrong default for agentic use at small scale "
        "— if the memory budget allows any move, spend it on "
        "precision before parameters, and treat published static-"
        "benchmark compression costs as a lower bound. For system "
        "builders: argument-level verification is not a free lunch; below "
        "a capability threshold it subtracts value, and failure "
        "composition (Figure 3) predicts where — selection- and "
        "synthesis-stage failures are out of its reach by construction. "
        "Mitigations that act on tool selection, or that repair instead "
        "of re-prompt, are the natural next experiments. For evaluators: "
        "single-shot leaderboard accuracy hides both the repetition "
        "collapse (Figure 2) and the failure-mode shifts that "
        "quantization induces; pass^k and per-stage attribution are cheap "
        "to add and change the conclusions."), first_indent=False)

    # ---- 7 Limitations ----
    heading(doc, "7  Limitations", 1)
    body(doc, (
        "The world is synthetic and single-domain, with 25 hand-written "
        "tasks; suite composition necessarily shapes the failure mix, and "
        "task-clustered bootstrap mitigates but cannot remove this. "
        "Per-cell Wilson intervals treat episodes as independent although "
        "they cluster by task, so those intervals are optimistic; all "
        "contrast inference respects the pairing. The 1.5B-Q4 floor (15%) "
        "leaves limited room to detect guardrail effects in that cell, a "
        "risk documented before the matrix ran. The grader is programmatic "
        "and was blind-validated on 42 episodes; post-fix agreement is "
        "in-sample, and residual grader error on unusual phrasings is "
        "possible. We tested one guardrail design; different error "
        "wording, budgets, or in-context repair examples may behave "
        "differently. Seeded sampling is reproducible on the same "
        "software and hardware, but bit-exact reproduction across GPUs is "
        "not guaranteed — which is why the raw logs, not the "
        "hardware, are the artifact of record. The extension matrix was "
        "not pre-registered and is reported as exploratory."),
        first_indent=False)

    # ---- 8 Conclusion ----
    heading(doc, "8  Conclusion", 1)
    body(doc, (
        "We measured, rather than assumed, the reliability of the "
        "language-model agents anyone can run for free. Quantization "
        "emerged as a first-order variable that changes how small agents "
        "fail, not just how often; the cheapest standard mitigation "
        "turned out to help only above a capability threshold and to "
        "hurt below it; and repetition exposed unreliability that "
        "single-shot scores conceal. Everything — harness, tasks, "
        "2,400 raw episode logs, analysis, and figures — is public "
        "and regenerates in CI, so every claim in this paper can be "
        "checked, and every future small model can be dropped into the "
        "same testbed in an afternoon."), first_indent=False)

    heading(doc, "Reproducibility statement", 1)
    body(doc, (
        "Code, task suite, raw logs for all 2,400 episodes, analysis "
        "scripts, and figure generation are available at "
        "https://github.com/DhruvilShah22/slm-agent-eval. A CI workflow "
        "regenerates the synthetic world (asserting determinism), "
        "re-validates the grader against the committed blind labels, and "
        "rebuilds every table and figure from the raw logs on each push. "
        "The full experiment reruns in under four GPU-hours on a free "
        "Kaggle tier, or on a consumer CPU. No paid APIs were used at any "
        "stage."), first_indent=False)

    # ---- References ----
    heading(doc, "References", 1)
    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph(f"[{i}] {ref}")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        for run in p.runs:
            run.font.size = Pt(9.5)

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
